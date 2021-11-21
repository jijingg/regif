#!/usr/bin/env python
# coding=utf-8
# Date   : 2017.7.28
# Author : jijing.guo (All rights reserved) 
# Email  : goco.v@163.com
# version: 0.10(realize RW/RO/WO/RWW/RWC/RWT feature)
#        : 1.00(code refactor, add RC/RWP/P(Protect) feature)
#        : 1.10(add word.docx generate feature 2018.3.27)
#        : 1.20(add keword FIFO_IF)
#        : 1.30(add ClockGate on RW/RWT/WO reg)
#        : 1.31(argparser supported)
#        : 1.32(add read Group feature)
#        : 1.33(add reset-macro)
#        : 1.8(add reset-macro)

import sys, os, json, shutil 
import logging, re
from typing import *

from docx.api import Document

from regif import RWC_RST_L1

def checklib():
    info = ""
    try:
        import xlrd
    except ImportError as e:
        info += str(e) + ", 'pip install xlrd' first!! \n"
    try:
        from docx import Document
    except ImportError as e:
        info += str(e) + ", 'pip install python_docx' first!! \n"
    if(info):
        log.error(info)
        exit(0)
    
checklib()

# logging 
class CustomFormatter(logging.Formatter):
    """Logging Formatter to add colors and count warning / errors"""
    grey     = "\x1b[38;1m"
    yellow   = "\x1b[33;1m"
    green    = "\x1b[32;1m"
    red      = "\x1b[31;1m"
    bold_red = "\x1b[31;1m"
    reset    = "\x1b[0m"
    format = "%(levelname)s - %(message)s"
    # format = "%(levelname)s - %(message)s (%(filename)s:%(lineno)d)"

    logging.SUCCESS = 25  # between WARNING and INFO
    logging.addLevelName(logging.SUCCESS, 'SUCCESS')
    FORMATS = {
        logging.DEBUG:    grey     + format + reset,
        logging.INFO:     grey     + format + reset,
        logging.WARNING:  yellow   + format + reset,
        logging.ERROR:    red      + format + reset,
        logging.SUCCESS:  green    + format + reset,
        logging.CRITICAL: bold_red + format + reset
    }

    def format(self, record):
        log_fmt = self.FORMATS.get(record.levelno)
        formatter = logging.Formatter(log_fmt)
        return formatter.format(record)

log = logging.getLogger("regif")
log.setLevel(logging.DEBUG)
ch = logging.StreamHandler()
ch.setFormatter(CustomFormatter())
log.addHandler(ch)
setattr(log, 'success', lambda message, *args: log._log(logging.SUCCESS, message, args))

# common
reg_format = [
    ("offset ", r"0x[\da-fA-F]+\Z", '"0x001C"'),
    ("regname", r"[a-zA-Z]\w*\Z", '"M_DMA_CFG"'),
    ("reg_des", r".*", 'some description is better '),
    ("width  ", r"[1-9]\d{1,2}\Z", '1~999 is better')]

field_format = [
    ("section", r"\[(\d+|\d+\:\d+)\]\Z", '"[0]" or "[15:7]" '),
    ("fd_name", r"[a-zA-Z]\w*\Z", 'must be one word '),
    ("RW     ", r"(RW|RC|BP|CP|RO|WO|RWW|RWC|RWT|RWP)\Z", "RW,RC,BP,CP,RO,WO,RWW,RWC,RWT,RWP only"),
    ("reset  ", r"(\d+'(h|H)[\da-fA-F]+\Z|\d+'(d|D)\d+\Z|\d+'(b|B)[01]+\Z|0\Z)", '''"0","3'd7", "12'hFFF", "2'b11"'''),
    ("wprotect", r"BITWISE|bitwise|\[(\d+|\d+\:\d+)\]\Z|\A\Z", "'[0]' '[12:8]' 'bitwise' 'BITWISE' or empty is ok"),
    ("locksig", r"[a-zA-Z]\w*\Z|\A\Z", "empty or one word be allowd"),
    ("fd_des ", r".*", ""),
    ("other  ", r".*", "")]

other_format = [
    (1, 1, "moudle ", r"[a-zA-Z]\w*", "first word need be a valid word"),
    (3, 5, "addr_wdith ", r"([2-9]\Z|[1-9]\d{1,2}\Z)", "2~999 is better"),
    (4, 5, "dw ", r"([4-9]\Z|[1-9]\d{1,2}\Z)", "4~999 is better"),
    (6, 5, "reg_pre ", r'\"[a-zA-Z_]\w*\"\Z|\"\"\Z', '"if_","reg_","" are better')]

# template 
ALWAYS_BEGIN = '''always @(posedge bmi_clk or negedge bmi_rstn)
    if(!bmi_rstn) begin '''
ALWAYS_BEGIN_MACRO="`ctrl_always_ff_begin(bmi_clk, bmi_rstn)"

RESET_MACRO="""
`ifndef __reset_macro_vh__
  `define __reset_macro_vh__
  `ifdef CTRL_RESET_ASYNC_LOW
    `define ctrl_always_ff_begin(clk,rst) always @(posedge clk or negedge rst) if(!rst) begin
  `elsif CTRL_RESET_ASYNC_HIGH
    `define ctrl_always_ff_begin(clk,rst) always @(posedge clk or posedge rst) if(rst) begin 
  `elsif CTRL_RESET_SYNC_LOW
    `define ctrl_always_ff_begin(clk,rst) always @(posedge clk) if(!rst) begin
  `elsif CTRL_RESET_SYNC_HIGH
    `define ctrl_always_ff_begin(clk,rst) always @(posedge clk) if(rst) begin
  `else             
    `define ctrl_always_ff_begin(clk,rst) always @(posedge clk or negedge rst) if(!rst) begin
  `endif   
`endif
"""

# common function
get_regif_dir = lambda : os.path.split(os.path.realpath(__file__))[0]

def trystr(value, ignorecase=0):
    _tmp = str(int(value)) if type(value) == float else \
        str(value) if type(value) == int else value
    return _tmp.lower() if ignorecase else _tmp

def tryint(value):
    s = trystr(value)
    return int(s)

def trybool(value, expect):
    return True if (trystr(value).upper() == expect) else False

def format_check(strs, line, row, info):
    if not re.match(info[1], strs):
        wstrs = '"' + strs + '"'
        return "Error: position {:3}{}: {:28} format error! Suggest like : {}\n" \
            .format(line + 10, chr(65 + row), wstrs, info[2])
    else:
        return ""

def title_check(bs):
    msg = ""
    for formats in other_format:
        line = formats[0]
        row = formats[1]
        info = formats[2:]
        msg += format_check(trystr(bs.cell_value(line, row)), line, row, info)
    return msg

def split2fixwidth(strlists, width):
    _out = "\n    "
    _len = 0
    for _str in strlists:
        _step = (len(_str) + 2)
        _len += _step
        if _len > width and (width - (_len - _step)) < 8:
            _len = 0
            _out += ("\n    " + _str + ", ")
        else:
            _out += (_str + ", ")
    return _out

def ispow2(x):
    return False if x & (x-1) else True

def groupSize(info):
    if info:
        gs = int(info)
        if(ispow2(gs)):
            return gs
        else:
            log.error("%d is not pow2, 32, 64, 128 is recommended, Fix at '5K' postion, then re-run again\n"%gs)
            exit()
    else:
        return 0

def getopt(args):
    withopt = [x.strip("opt=") for x in args if "opt=" in x]
    if withopt:
        return withopt[0]
    else :
        return ""

def intger(s: str):
    if(s.startswith("0x")):
        return int(s, 16)
    elif(s.startswith("0o")):
        return int(s, 8)
    elif(s.startswith("0b")):
        return int(s, 2)
    else :
        return int(s, 10)

def dumpfile(path, content: str):
    with open(path, "w") as f:
        f.write(content)
    log.success("{} generate done!".format(path))

# object 
class Wire:
    dc_name = "wire"
    def __init__(self, name, dw = 1):
        self.name = name 
        self.dw = dw 
        self.align = 20

    def __str__(self):
        ndw = "" if(self.dw == 1) else "[{}:0]".format(self.dw-1)
        ndw = ndw.center(10)
        nname = self.name.ljust(self.align)
        dcname = self.dc_name.ljust(10)
        return  "{} {} {}".format(dcname, ndw, nname)
                                                 
    def __repr__(self):
        return self.__str__()

class Input(Wire):
    dc_name = "input"

class Output(Wire):
    dc_name = "output"
 
class OutputReg(Wire):
    dc_name = "output reg"

bmiport = [
    Input("bmi_clk"),
    Input("bmi_rstn"),
    Input("bmi_rd"),
    Input("bmi_wr"),
    Input("bmi_addr", 16),
    Input("bmi_wdata", 32),
    Output("bmi_rdata", 32),
    Output("bmi_rdvld"),
]

class Declars:
    jdot = ";\n"
    def __init__(self, signals):
        self.maxwidth = max([len(s.name) for s in signals])
        self.signals = map(self.align, signals)

    def align(self, s):
        s.align = self.maxwidth + 2
        return s

    def _joind(self):
        return self.jdot.join(map(str, self.signals))

    def __repr__(self):
        return self.__str__()

    def __str__(self):
        return self._joind() + ";"

class IODeclars(Declars):
    jdot = ",\n"
    def __str__(self):
        return self._joind()

class InstPort():
    def __init__(self, pname, wname, io, dw = 1):
        self.port = pname
        self.wire = wname
        self.io = io
        self.pgap = 30
        self.wgap = 30
        self.dw = dw
        self.islast = False
                                          
    def __str__(self):
        pname = self.port.ljust(self.pgap)
        dot = " " if self.islast else ","
        io = self.io
        ndw = "" if(self.dw == 1) else "[{}:0]".format(self.dw-1)
        wname = (self.wire+ndw).ljust(self.wgap+10)
        return  "    .{} ( {} ){}//{}".format(pname, wname, dot, io)

class InstPorts():
    def __init__(self, ports):
        self.pmax = max([len(p.port) for p in ports])
        self.wmax = max([len(p.wire) for p in ports])
        self.ports = list(map(self.align, ports))
        self.ports[-1].islast = True

    def align(self, s):
        s.pgap = self.pmax + 2
        s.wgap = self.wmax + 2
        return s

    def __str__(self):
        return "\n".join(map(str, self.ports))
    def __repr__(self):
        return self.__str__()  

class Field:
    "sec/name/acc/reset/wp/lock/doc"
    def __init__(self, fieldjs):
        'deep copy is very important unless it will silence change fieldjs dict'
        import copy
        js = copy.deepcopy(fieldjs)
        js["sec"] = js2sec(js["sec"])
        self.raw_reset = js["reset"]
        js["reset"] = resetvalue(self.raw_reset)
        self.__dict__.update(js)
        self.max_val = 1 << self.sec.width

    def __repr__(self) -> str:
        return self.__str__()
    def __str__(self) -> str:
        return "filed:{}{}".format(self.name, self.sec)

    def write_str(self) -> str:
        return ""
    def read_str(self) -> str:
        return ""
    def io(self):
        iolist = []
        return iolist
    def dec(self):
        declist = []
        return declist
    def reset_check(self, pre):
        msg = ""
        cm = re.match("(\d+)'(d|h|b)([\da-fA-F]+)", self.raw_reset)
        if cm and self.sec.width != int(cm.group(1)):
            msg =  "{pre}field {name:10} width {sec.width} don't match it reset value \"{raw_reset}\" \n".format(pre = pre, **self.__dict__)
        if self.reset > (self.max_val - 1):
            max = 1 << self.sec.width - 1
            msg += "{pre}field {name:10} reset value {raw_reset} exceed max value \"{max}\"\n".format(pre = pre, max = max, **self.__dict__)
        return msg

"RW/RO/WO/RWP/RWT/RC/WC/RWC/RWW"

class Reg:
    "name/offset/doc"
    def __init__(self, name, body, dw = 32):
        self.name = name.strip()
        self.offset = body["offset"]
        self.doc = body["doc"]
        self.fields = self.getfields(body["fields"]) 
        self.signamemax = max([len(fd.name) for fd in self.fields])
        self.io = self.getiosignals()
        self.align = 30
        self.dw = dw

    def getfields(self, fdjs):
        ret = [Field(fd) for fd in fdjs]
        ret.reverse()
        return ret

    def upsigalign(self, l):
        for fd in self.fields:
            fd.align = l

    def getiosignals(self):
        pass

    def __repr__(self) -> str:
        return self.__str__()
    def __str__(self) -> str:
        return "reg:{} {} [{}]".format(self.name, self.offset, "|".join([str(s) for s in self.fields]))

    def write_str(self) -> str:
        return ""
    def read_str(self) -> str:
        return ""
    def io_str(self) -> str:
        return ""

    def dynamic_check(self) :
        pre = "Error: reg {name:{align}} ".format(**self.__dict__)
        msg = self.sec_continue_check(pre)
        fdrstck = [fd.reset_check(pre) for fd in self.fields]
        nfd = [fd for fd in fdrstck if fd]
        msg += "\n".join(nfd)
        return msg

    def sec_continue_check(self, pre):
        msg = ""
        s0 = self.fields[0]
        sm = self.fields[-1]
        render_dict = self.__dict__
        render_dict.update(locals())
        if s0.sec.lsb != 0:
            msg += "{pre}first section {s0} not begin with 0 \n".format(**render_dict)
        if sm.sec.msb != self.dw - 1:
            msg += "{pre}last section {sm} not end with {dw}\n".format(**render_dict)
        if len(self.fields) > 1:
            for i in range(1, len(self.fields)):
                sp = self.fields[i-1]
                sn = self.fields[i]
                render_dict.update(locals())
                if(sp.sec.msb + 1) != sn.sec.lsb:
                    msg += "{pre}section {sn} {sp} not continuity \n".format(**render_dict)
        return msg

    def proteck_check(self):
        msg = ""
        return msg
# template 

VFILE = """
module {name}_regif({io}); 
{dec}
endmodule
"""

WFILE = """
module {name}_regif_wrap({io});
{dec}
{inst}
{lock}
endmodule
"""

INST = """
{name}_regif u_{name}_regif(
    {instports}
);
"""

class RegIf:
    def __init__(self, regs, name, cg: bool, grpsize: int, rstmacro: int):
        self.name = name
        self.regs = regs
        self.withcg = cg
        self.grpsize = grpsize
        self.withrstmacro = rstmacro
        self.locksigs = self.getlocksigs()
        self.withlock = self.locksigs != []

    def getlocksigs(self):
        locksig = []
        for reg in self.regs:
            for field in reg.fields:
                if field.lock.strip():
                    locksig.append(field.lock)
        return locksig

    def io_str(self):
        return "j"
    def dec_str(self):
        pass
    def write_str(self):
        pass
    def read_str(self):
        pass
    def wrap_io_str(self):
        pass
    def wrap_lock_str(self):
        pass
    def wrap_dec_str(self):
        pass
    def wrap_inst_str(self):
        return INST.format(**self.__dict__)

    def vfile(self):
        iopart = self.io_str()
        declare = self.dec_str()
        return VFILE.format(io = iopart, dec = declare, **self.__dict__)

    def wvfile(self):
        return "regif_wrap.v"
        inst = self.wrap_inst_str()
        return WFILE.format(**self.__dict__)

    def dump(self, path = "./"):
        vpth = os.path.join(path, self.name + "_regif.v")
        vwpth = os.path.join(path, self.name + "_regif_wrap.v")
        dumpfile(vpth, self.vfile())
        if self.withlock:
            dumpfile(vwpth, self.wvfile())

def js2sec(js):
    return Section(js["msb"], js["lsb"])

def str2sec(s):
    multi_bit = re.match(r"\[(?P<b>\d+):(?P<s>\d+)\]", s)
    single_bit = re.match(r"\[(?P<b>\d+)\]", s)
    if multi_bit:
        msb = int(multi_bit.group(1))
        lsb = int(multi_bit.group(2))
        return Section(msb, lsb)
    elif single_bit:
        msb = int(single_bit.group(1))
        return Section(msb, msb)
    else:
        raise "bit section error"

def resetvalue(s):
    decm = re.match("(\d+)'d(\d+)", s)
    binm = re.match("(\d+)'b([01]+)", s)
    hexm = re.match("(\d+)'h([\da-f]+)", s)
    if s.isdigit():
        rel_val = int(s)
    elif decm:
        rel_val = int(decm.group(2))
    elif binm:
        rel_val = int(binm.group(2), 2)
    elif hexm:
        rel_val = int(hexm.group(2), 16)
    else:
        raise "An except error, the reset value not match 'd 'b 'h 0 type"
    return rel_val

class Section:
    def __init__(self, msb, lsb):
        self.msb = msb
        self.lsb = lsb
        self.width = msb - lsb + 1
        self.single = msb == lsb

    def __str__(self):
        return "" if(self.msb == self.lsb)  else "[{}:{}]".format(self.msb, self.lsb)
    def __repr__(self):
        return self.__str__()  
    def doc(self):
        return "[{}]".format(self.msb) if(self.single)  else "[{}:{}]".format(self.msb, self.lsb)

from enum import Enum
class Acc(Enum):
    RW = 1
    WO = 2
    RO = 3
    RC = 4
    WC = 5
    WT = 6
    WP = 7

# excel load and parser
class MSDoc():
    def __init__(self, js):
        from docx import Document
        self.jsregs = js
        self.doc = Document()

    def filldoc(self):
        from docx.shared import RGBColor
        from docx.oxml.ns import qn
        for name in self.jsregs:
            body = self.jsregs[name]
            self.doc.styles['Normal'].font.name = u'Times New Roman'
            self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Times New Roman')
            self.doc.add_heading(u'{} ({})'.format(name, body["offset"]), 2)
            self.doc.add_paragraph(u"Offset: {}".format(body["offset"]))
            self.doc.add_paragraph(u"Description: {}".format(body["doc"]))
            reg_table = self.creat_reg_table()
            self.filltable(reg_table, body["fields"])

    def creat_reg_table(self):
        from docx.shared import Pt
        from docx.shared import Cm
        from docx.shared import Length
        table_style = self.doc.styles["Table Grid"]
        table_style.font.size = Pt(10.5)
        table = self.doc.add_table(rows=1, cols=0, style=table_style)
        aval_width = Length(self.doc._block_width).cm
        table.add_column(Cm(aval_width * 0.1))
        table.add_column(Cm(aval_width * 0.15))
        table.add_column(Cm(aval_width * 0.1))
        table.add_column(Cm(aval_width * 0.15))
        table.add_column(Cm(aval_width * 0.45))
        table.cell(0, 0).text = u"Width"
        table.cell(0, 1).text = u"RegName"
        table.cell(0, 2).text = u"AccType"
        table.cell(0, 3).text = u"ResetValue"
        table.cell(0, 4).text = u"Description"
        return table

    def filltable(self, table, fields):
        for field in fields:
            row_cells = table.add_row().cells
            reserved = True if field["doc"] == "RESERVED" else False
            row_cells[0].text = js2sec(field["sec"]).doc()
            row_cells[1].text = field["name"]
            row_cells[2].text = "N/A" if reserved else trystr(field["acc"])
            row_cells[3].text = "N/A" if reserved else trystr(field["reset"])
            row_cells[4].text = trystr(field["doc"])

    def dump(self, name, dir = "./"):
        self.doc.add_heading(u'Register Interface', 1)
        self.filldoc()
        path = os.path.join(dir, u"{}_Register_Manual.docx".format(name))
        self.doc.save(path)
        log.success("{} generate done!".format(path))

# Global SingleTone Config object
def singleton(cls):
    instance = []
    def single(*args, **kwargs):
        if len(instance)==0:
            instance.append(cls(*args, **kwargs))
        return instance[0]
    return single

@singleton
class GlobalConfig:
    def __init__(self):
        self.aw : int = 32
        self.dw : int = 32
        self.withcg: bool = False
        self.withrstmacro: bool = False
        self.groupsize: int = 0
        self.lock_pre: str = ""
        self.reg_pre : str = ""
        self.hwclc : str = "_hw_clc"
        self.hwset : str = "_hw_set"
        self.hwsetval : str = "_hw_set_val"
        self.outdir = "./"
        self.regaw: int = 32
        self.sigaw: int = 32
        self.always_ff_begin = ALWAYS_BEGIN_MACRO if(self.withrstmacro) else ALWAYS_BEGIN
        self.macros = RESET_MACRO if(self.withrstmacro) else ""
        self.unlocktype = ("RO", "RC", "RWT", "RWP", "CP", "BP")

    @property
    def awmsb(self) : self.aw - 1
    @property
    def dwmsb(self) : self.dw - 1

    def dict(self):
        return self.__dict__
    
class XLSParser():
    _regs = []
    _js = {}
    def __init__(self, bs):
        self._static_check(bs)
        self.getjson()
        self.name         = re.match(other_format[0][3], bs.cell(2, 1).value).group()
        self.aw           = tryint(bs.cell(3, 5).value)
        self.dw           = tryint(bs.cell(4, 5).value)
        self.withcg       = trybool(bs.cell(3, 10).value, "YES")
        self.reg_pre      = trystr(bs.cell(6, 5).value).strip('"')
        self.groupsize    = groupSize(bs.cell(4, 10).value)
        self.withrstmacro = bs.cell(5,10).value.upper() == "YES"
        self.regs         = json2reg(self._js)
        self.regif        = RegIf(self.regs, self.name, self.withcg, self.groupsize, self.macros)
        self._dynamic_check()

    def __repr__(self) -> str:
        return "\n".join([str(s) for s in self.regs])

    def getjson(self):
        def sec(s):
            ns = str2sec(s)
            return {"msb" : ns.msb, "lsb" : ns.lsb}

        def field(t):
            return {"sec": sec(t[0]), "name":t[1], "acc" : t[2],  "reset" : trystr(t[3]), "wp": t[4], "lock": t[5], "doc":t[6]}

        def reg(x):
            regname = x[1]
            fields = [field(t) for t in x[4]]
            return { regname : {
                "offset" : intger(x[0]),
                "doc" : x[2],
                "fields" : fields}
            }

        def tryupdate(x):
            key = list(x.keys())[0]
            if key in self._js:
                log.error("{} already exits".foramt(key))
                exit(2)
            self._js.update(x)

        [tryupdate(reg(x)) for x in self._regs]

    def getregs(self, bs):
        nrows = bs.nrows
        ncols = bs.ncols
        address_idx = 0
        for i in range(9, nrows):
            row_data = bs.row_values(i)
            if row_data[0] != u'':
                self._regs.append(row_data[0:4])
                self._regs[address_idx].append([])
                self._regs[address_idx][4].append(row_data[4:ncols])
                address_idx += 1
            else:
                self._regs[address_idx - 1][4].append(row_data[4:ncols])

    def _dynamic_check(self):
        log.info("Dynamic checking ..... ")
        msg = ""
        for reg in self.regs:
            msg += reg.dynamic_check()
        if msg:
            msg += "\nDynamic check Fail, fixed all those format Error first, then regenerate again.\n"
            log.warning("\n\n" + msg)
            exit(0)
        else:
            log.success("Dynamic check Pass!")
        return

    def _static_check(self, bs):
        self.getregs(bs)
        log.info("Static checking ..... ")
        msg = title_check(bs)
        msg += static_check(self._regs)
        if msg:
            msg += "\nStatic check Fail, fixed all those format Error first, then regenerate again.\n"
            log.warning("\n\n" + msg)
            exit(0)
        else:
            log.success("Static check Pass!")
        return

    def dumpjson(self):
        content = json.dumps(self._js, sort_keys=False, indent=4)
        jspth = os.path.join(self.outdir, "{}.json".format(self.name))
        dumpfile(jspth, content)

    def dumpdoc(self):
        doc = MSDoc(self._js)
        doc.dump(self.name, self.outdir)

    def dumpverilog(self):
        self.regif.dump()

    def dump(self, args):
        self.dumpverilog()
        if( "--doc" in args):
            self.dumpdoc()
        if( "--json" in args):
            self.dumpjson()

def json2reg(js):
    regs = [Reg(name, js[name]) for name in js]
    regnamemax = max(map(lambda x: len(x.name), regs))
    # signamemax = max(map(lambda x: x.signamemax, regs))
    def upalign(reg):
        reg.align = regnamemax
        # reg.upsigalign(signamemax)
        return reg
    return list(map(upalign, regs))

def static_check(regs):
    msg = ""
    line = 0
    regs_unq_msg, addrs_unq_msg, fields_unq_msg = ("",) * 3
    regs_unq, addrs_unq, fields_unq = ([],) * 3
    for reg in regs:
        if reg[1] in addrs_unq:
            addrs_unq_msg += "Error: position {:3}{}: Regname {:20} already exist\n".format(line + 10, chr(65 + 1), reg[1])
        if reg[0] in regs_unq:
            regs_unq_msg += "Error: position {:3}{}: Addrs   {:20} already exist\n".format(line + 10, chr(65 + 0), reg[0])
        addrs_unq.append(reg[1])
        regs_unq.append(reg[0])
        for i in range(len(reg_format)):
            strs = trystr(reg[i])
            msg += format_check(strs, line, i, reg_format[i])
        for field in reg[4]:
            if field[1] in fields_unq and field[1] != "RESERVED":
                fields_unq_msg += "Error: position {:3}{}: Field   {:20} already exist\n".format(line + 10, chr(65 + 4), field[1])
            fields_unq.append(field[1])
            for i in range(len(field_format)):
                strs = trystr(field[i])
                msg += format_check(strs, line, i + 4, field_format[i])
            line += 1
    return msg + regs_unq_msg + addrs_unq_msg + fields_unq_msg


# argvs process
def checkargs(args, info):
    if(len(args) == 0):
        log.error("misuage, '{}' try again".format(info))
        exit(2)

def creatxls(name):
    template = os.path.join(get_regif_dir(), "regif_template.xls")
    dest = "%s_regif.xls" % name
    if os.path.exists(dest):
        log.error("%s already exists !!!" % dest)
        exit(1)
    shutil.copy(template, dest)
    log.info("RegIf excel template %s created\n" % dest)

def task_regif(args):
    import xlrd
    checkargs(args, "regif --init youmodule.xls")
    opt = getopt(args)
    xls = args[0]
    if (not os.path.exists(xls)):
        log.error("%s not found, check please, 'regif -h' for help" % xls)
        exit(2)
    wb = xlrd.open_workbook(xls)
    bs = wb.sheet_by_index(0)
    xlsreg = XLSParser(bs)
    xlsreg.dump(args)

def task_creat(args):
    checkargs(args, "regif --init youmodule.xls")
    creatxls(args[1])

def task_ui(args):
    help = """
    Welcome regif (register interface generator tools)

    * regif --init mymodule            //creat template
    * regif mymodule.xls               //generate .v
    * regif mymodule.xls --doc         //generate .v .doc
    * regif mymodule.xls --doc --json  //generate .v .doc .json
    * regif mymodule.xls --html        //generate .v .html(not ready)
    * regif mymodule.xls --pdf         //generate .v .pdf(not ready)
    """
    print(help)

task_entry = {
    "--init" : task_creat, 
    "xls"    : task_regif,
    "--help" : task_ui, 
    "-h"     : task_ui, 
}

def main(argvs):
    noargs = not argvs
    if(noargs):
        task_ui([])
    else:
        command = argvs[0]
        args    = argvs
        command = "xls" if ".xls" in command else command
        task = task_entry.get(command, task_ui)
        task(args)

if __name__ == "__main__":
    try: 
        argvs = sys.argv[1:]
        main(argvs)
    except Exception as e:
        log.error("regif crush, report to https://github.com/jijingg/regif/issues")
        raise e
