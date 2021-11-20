#!/usr/bin/env python
# coding=utf-8
# Date   : 2017.7.28
# Author : jijing.guo (All rights reserved) 
# Email  : goco.v@163.com
# version: 0.10(realize RW/RO/WO/RWW/RWC/RWT feature)
#        : 1.00(code refactor, add RC/RWP/P(Protect) feature)
#        : 1.10(add word.docx generate feature 2018.3.27)
#        : 1.20(add keword FIFO_IF)
#        : 1.30(add ClockGate on RW/RWT/WO reg)
#        : 1.31(argparser supported)
#        : 1.32(add read Group feature)
#        : 1.33(add reset-macro)


def checklib():
    info = ""
    try:
        import xlrd
    except ImportError as e:
        info += str(e) + ", 'pip install xlrd' first!! \n"
    try:
        from docx import Document
    except ImportError as e:
        info += str(e) + ", 'pip install python_docx' first!! \n"
    if(info):
        info = "Error: \n" + info 
        print(info)
        exit(0)
    
checklib()
import xlrd
from docx import Document

import re
import json
import sys, os
import argparse

# python 2.x defualt encode use ascii
# python 3.x defualt encode use unicode
if sys.version_info < (3, 0):
    reload(sys)
    sys.setdefaultencoding("utf-8")  # for chinese charactor encode

reg_format = [
    ("offset ", r"0x[\da-fA-F]+\Z", '"0x001C"'),
    ("regname", r"[a-zA-Z]\w*\Z", '"M_DMA_CFG"'),
    ("reg_des", r".*", 'some description is better '),
    ("width  ", r"[1-9]\d{1,2}\Z", '1~999 is better')]

field_format = [
    ("section", r"\[(\d+|\d+\:\d+)\]\Z", '"[0]" or "[15:7]" '),
    ("fd_name", r"[a-zA-Z]\w*\Z", 'must be one word '),
    ("RW     ", r"(RW|RC|BP|CP|RO|WO|RWW|RWC|RWT|RWP)\Z", "RW,RC,BP,CP,RO,WO,RWW,RWC,RWT,RWP only"),
    ("reset  ", r"(\d+'(h|H)[\da-fA-F]+\Z|\d+'(d|D)\d+\Z|\d+'(b|B)[01]+\Z|0\Z)", '''"0","3'd7", "12'hFFF", "2'b11"'''),
    ("wprotect", r"BITWISE|bitwise|\[(\d+|\d+\:\d+)\]\Z|\A\Z", "'[0]' '[12:8]' 'bitwise' 'BITWISE' or empty is ok"),
    ("locksig", r"[a-zA-Z]\w*\Z|\A\Z", "empty or one word be allowd"),
    ("fd_des ", r".*", ""),
    ("other  ", r".*", "")]

other_format = [
    (1, 1, "moudle ", r"[a-zA-Z]\w*", "first word need be a valid word"),
    (3, 5, "addr_wdith ", r"([2-9]\Z|[1-9]\d{1,2}\Z)", "2~999 is better"),
    (4, 5, "data_width ", r"([4-9]\Z|[1-9]\d{1,2}\Z)", "4~999 is better"),
    (6, 5, "reg_pre ", r'\"[a-zA-Z_]\w*\"\Z|\"\"\Z', '"if_","reg_","" are better')]
# field joined check
INPUT = "    input     {bw:11}{pn:{gfd_align}} ,\n"
OUTPUT = "    output    {bw:11}{pn:{gfd_align}} ,\n"
WIRE = "wire      {bw:11}{pn:{gfd_align}} ;\n"
REG = "reg       {bw:11}{pn:{gfd_align}} ;\n"
IPORTS = "    .{pn:{gfd_align}} ({pn2:{gfd_align}} ),//i\n"
OPORTS = "    .{if_name:{gfd_align}} ({pn2:{gfd_align}} ),//o\n"
FIFOMAP = "//assign  {pn:20} = fifo_rdata{section:8} ;\n"
ALWAYS_BEGIN = '''always @(posedge bmi_clk or negedge bmi_rstn)
    if(!bmi_rstn) begin '''
ALWAYS_BEGIN_MACRO="`ctrl_always_ff_begin(bmi_clk, bmi_rstn)"
DEC_PORTS = \
    '''    //Inputs
    bmi_clk, bmi_rstn, bmi_rd, bmi_wr, bmi_addr, bmi_wdata, {}{}\n
    //Outputs
    bmi_rdata, bmi_rdvld{}'''

FIFO_PUSH_L4 = \
    '''//-------------------------------------------------------------------
//{addr} : {reg_name:{reg_align}} {if_Comment}
wire   {reg_name}_push ;
assign {reg_name}_push = bmi_wr && (bmi_addr==`{reg_name});
{fifo_filed_assign}\n'''

W_BLOCKS_L4 = \
    '''//-------------------------------------------------------------------
//{addr} : {reg_name:{reg_align}} {if_Comment}
{rw_cg_always}{w_always}{rc_part}{cdc_part}'''
RC_ALWYAS_L0 = \
    '''
{always_ff_begin}
        {if_name}{bitsel} <= 1'b0;
    end else if({name}{hw}{syncd}{bitsel}) begin 
        {if_name}{bitsel} <= 1'b1;//RC
    end else if(bmi_wr && (bmi_addr==`{reg_name}) && bmi_wdata{bmibitsel}) begin 
        {if_name}{bitsel} <= 1'b0;//RC
    end  
'''
W_RST_L1 = \
    "\n        {if_name:{fd_align}}  <= {rstval:14}  ;"
W_CG_ALWAYS_L3 = \
    '''{always_ff_begin}{rw_rst}
    end else if(bmi_wr && (bmi_addr==`{reg_name})) begin{rw_wr}{rwt_wr}{protect_wr}
    end 
'''
W_ALWAYS_L3 = \
    '''{always_ff_begin}{field_rst}
    end else if(bmi_wr && (bmi_addr==`{reg_name})) begin{normal_wr}{rwt_wr}{protect_wr}{else_part}
    end 
'''
BP_RWT_L0 = \
    "\n        if(bmi_wdata{bmibitselp:4}) {if_name}{bitsel:4} <= ~{if_name}{bitsel:4};//{types}"
BP_L0 = \
    "\n        if(bmi_wdata{bmibitselp:4}) {if_name}{bitsel:4} <= bmi_wdata{bmibitsel:4};//{types}"
W_L1 = \
    "\n        {tab}{if_name:{fd_align}}  <= bmi_wdata{section:7};//{type}"
RWT_L1 = \
    "\n        {tab}{if_name:{fd_align}}  <= bmi_wdata{section}? ~{if_name}:{if_name};//{type}"
RWC_RST_L1 = \
    '''\n        if({name}{hw}{syncd}) begin
            {if_name:{fd_align}}  <= {rstval:6}  ;//RWC
        end '''
RWW_SET_L1 = \
    '''\n        if({name}{seten}) begin
            {if_name:{fd_align}}  <= {name}{setval}   ;//RWW
        end '''
ELSE_PART_L2 = \
    "\n    end else begin{fieldrwp}{fieldrwc}{fieldrww}"
W_PROTECTE_L2 = \
    '''\n        //CP(condtion protection) write
        if(bmi_wdata{p_section} == {w_valid_val}) begin {cprotects}
        end'''
CDC_RST_L1 = ''' 
        {name}{hw}_q1{bitsel:3} <=  1'b0 ;
        {name}{hw}_q2{bitsel:3} <=  1'b0 ;'''
CDC_FF_L1 = '''
        {name}{hw}_q1{bitsel:3} <=  {name}{hw}{bitsel:3}    ;
        {name}{hw}_q2{bitsel:3} <=  {name}{hw}_q1{bitsel:3} ;'''
CDC_PLS_L1 = '''
assign {name}{hw}{syncd}{bitsel:3} = {name}{hw}_q1{bitsel} && ~{name}{hw}_q2{bitsel} ;'''
R_L2 = \
    '''\n            `{reg_name:{reg_name_align}} : begin {bmi}_rdvld <=1'b1; {bmi}_rdata <=  {{{_read}}} ;end'''
CDC_PART_L2 = \
    '''
//cdc porecess, sync signal from HW to Bmi_clk domain
//only support RC/RWC, not support RWW
{always_ff_begin}{cdc_sync_rsts}
    end else begin {cdc_sync_ffs}
    end

//we suppose the signals from HW are pluse which width cover one bmi_clk cycle at least {cdc_sync_plss}
'''
R_ALWAYS = \
    '''//register bus read interface 
reg    [{data_width}-1 :0]     {bmi}_rdata      ;
reg                  {bmi}_rdvld      ;
{always_ff_begin}
        {bmi}_rdata <= {data_width}'b0 ;
        {bmi}_rdvld <= 1'b0  ;
    end else if(bmi_rd && ~bmi_wr) begin
        case(bmi_addr){reads}
            default : begin {bmi}_rdvld <= 1'b0; {bmi}_rdata <= {data_width}'b0 ;end
        endcase
    end else begin
        {bmi}_rdvld <= 1'b0  ;
    end
'''
RG_RD = \
    "\n        if     (grp{i}_rdvld) begin bmi_rdvld <= grp{i}_rdvld; bmi_rdata <= grp{i}_rdata; end"
RG_RDE = \
    "\n        else if(grp{i}_rdvld) begin bmi_rdvld <= grp{i}_rdvld; bmi_rdata <= grp{i}_rdata; end"
RG_ALWAYS = \
     '''//register bus read interface 
reg    [{data_width}-1 :0]     bmi_rdata      ;
reg                  bmi_rdvld      ;
reg                  rd_dly1        ;
{always_ff_begin}
        rd_dly1   <= 1'b0  ;
    end else begin
        rd_dly1   <= bmi_rd && ~bmi_wr ;
    end
        
{always_ff_begin}
        bmi_rdata <= {data_width}'b0 ;
        bmi_rdvld <= 1'b0  ;
    end else if(rd_dly1) begin{grpreads}
    end else begin
        bmi_rdvld <= 1'b0  ;
    end
'''       
CG_IO = '''\n    input                bmi_clk_wr_cg     ,
//Please instance manually outside like: 
//gate_cell  u_RW_cg_cell(.CLK(bmi_clk), .TSE(test_mode), .E(bmi_wr), .ECK(bmi_clk_wr_cg));'''
BMI_IO = \
    '''    input                bmi_clk           ,{}
    input                bmi_rstn          ,
    input                bmi_rd            ,
    input                bmi_wr            ,
    input    [{addr_widthp1:2}:0]      bmi_addr          ,
    input    [{data_widthp1:2}:0]      bmi_wdata         ,
    output   [{data_widthp1:2}:0]      bmi_rdata         ,
    output               bmi_rdvld         ,
'''
LOCK_EXPRESS_L1 = \
    "        {if_name:{fd_align}}  <= {lock_pre}{if_name:{fd_align}} ;\n"
LOCK_ALWAYS_L2 = \
    '''
always @(posedge hw_clk or negedge hw_rstn)
    if(!hw_rstn) begin{lock_rsts}    
    end else if({lock_en}) begin
{lock_express}    end
'''
RESET_MACRO="""
`ifndef __reset_macro_vh__
  `define __reset_macro_vh__
  `ifdef CTRL_RESET_ASYNC_LOW
    `define ctrl_always_ff_begin(clk,rst) always @(posedge clk or negedge rst) if(!rst) begin
  `elsif CTRL_RESET_ASYNC_HIGH
    `define ctrl_always_ff_begin(clk,rst) always @(posedge clk or posedge rst) if(rst) begin 
  `elsif CTRL_RESET_SYNC_LOW
    `define ctrl_always_ff_begin(clk,rst) always @(posedge clk) if(!rst) begin
  `elsif CTRL_RESET_SYNC_HIGH
    `define ctrl_always_ff_begin(clk,rst) always @(posedge clk) if(rst) begin
  `else             
    `define ctrl_always_ff_begin(clk,rst) always @(posedge clk or negedge rst) if(!rst) begin
  `endif   
`endif
"""
VFILE = \
    '''{defines}{macros}module {modulename}_regif(
{bmi_io}{iodeclars});
{vdeclars}
{write_part}
{read_part}
endmodule
{undefines}
'''
CG_PORT = "\n    .bmi_clk_wr_cg          (bmi_clk_wr_cg           ),//i"
WFILE = \
    '''module {modulename}_regif_wrap(
    input                hw_clk            ,
    input                hw_rstn           ,
{bmi_io}{lock_io}{iodeclars});
//======================================================
wire   [{data_widthp1:2}:0]        bmi_rdata         ;
wire                 bmi_rdvld         ;
{wdeclars}
//======================================================
{modulename}_regif u_{modulename}_regif(
    .bmi_clk                (bmi_clk                 ),//i{cg_port}
    .bmi_rstn               (bmi_rstn                ),//i
    .bmi_rd                 (bmi_rd                  ),//i
    .bmi_wr                 (bmi_wr                  ),//i
    .bmi_addr               (bmi_addr                ),//i
    .bmi_wdata              (bmi_wdata               ),//i
    .bmi_rdata              (bmi_rdata               ),//o
    .bmi_rdvld              (bmi_rdvld               ),//o      
{instports}
);
//======================================================
{lock_part}
endmodule 
'''

class Wire:
    dc_name = "wire"
    def __init__(self, name, dw = 1):
        self.name = name 
        self.dw = dw 
        self.align = 20

    def __str__(self):
        ndw = "" if(self.dw == 1) else "[{}:0]".format(self.dw-1)
        ndw = ndw.center(10)
        nname = self.name.ljust(self.align)
        dcname = self.dc_name.ljust(10)
        return  "{} {} {}".format(dcname, ndw, nname)
                                                 
    def __repr__(self):
        return self.__str__()

class Input(Wire):
    dc_name = "input"

class Output(Wire):
    dc_name = "output"
 
class OutputReg(Wire):
    dc_name = "output reg"

bmiport = [
    Input("bmi_clk"),
    Input("bmi_rstn"),
    Input("bmi_rd"),
    Input("bmi_wr"),
    Input("bmi_addr", 16),
    Input("bmi_wdata", 32),
    Output("bmi_rdata", 32),
    Output("bmi_rdvld"),
]

class Declars:
    jp = ";\n"
    def __init__(self, signals):
        self.maxwidth = max([len(s.name) for s in signals])
        self.signals = map(self.align, signals)

    def align(self, s):
        s.align = self.maxwidth + 2
        return s

    def _joind(self):
        return self.jp.join(map(str, self.signals))

    def __repr__(self):
        return self.__str__()

    def __str__(self):
        return self._joind() + ";"

class IODeclars(Declars):
    jp = ",\n"
    def __str__(self):
        return self._joind()
        # return super().__str__() + ";"     

class InstPort():
    def __init__(self, pname, wname, io, dw = 1):
        self.port = pname
        self.wire = wname
        self.io = io
        self.pgap = 30
        self.wgap = 30
        self.dw = dw
        self.islast = False
                                          
    def __str__(self):
        pname = self.port.ljust(self.pgap)
        dot = " " if self.islast else ","
        io = self.io
        ndw = "" if(self.dw == 1) else "[{}:0]".format(self.dw-1)
        wname = (self.wire+ndw).ljust(self.wgap+10)
        return  "    .{} ( {} ){}//{}".format(pname, wname, dot, io)

class InstPorts():
    def __init__(self, ports):
        self.pmax = max([len(p.port) for p in ports])
        self.wmax = max([len(p.wire) for p in ports])
        self.ports = list(map(self.align, ports))
        self.ports[-1].islast = True

    def align(self, s):
        s.pgap = self.pmax + 2
        s.wgap = self.wmax + 2
        return s

    def __str__(self):
        return "\n".join(map(str, self.ports))

    def __repr__(self):
        return self.__str__()  

class Reg:
    def __init__(self, reg, regifdict):
        self.__dict__.update(regifdict)
        self.reg_name = reg[1]
        self.addr = reg[0]
        self.rtl_free = True if "RTL_FREE" in reg[2] else False
        self.fifo_if = True if "FIFO_IF" in reg[2] else False
        self.addr = self.addr_width + "'h" + reg[0][2:]
        self.fields = RegFields(reg[4], self.__dict__)
        # self.fields= RegFields(reg[4],regifdict,self.name,self.addr)

    def dynamic_check(self):
        return self.fields.dynamic_check()

    @property
    def write(self):
        if (self.fifo_if):
            return self.fields.fifo_push_declar()
        return self.fields.w_block()

    def read(self, bmi="bmi"):
        return self.fields.read_express(bmi)

    @property
    def iodeclars(self):
        return self.fields.declars()[0]

    @property
    def vdeclars(self):
        return self.fields.declars()[1]

    @property
    def wdeclars(self):
        return self.fields.declars()[2]

    @property
    def instports(self):
        return self.fields.declars()[3]

    @property
    def defines(self):
        return "`define   {reg_name:{reg_align}}    {addr} \n".format(**self.__dict__)

    @property
    def undefines(self):
        return "`undef   {reg_name:{reg_align}} \n".format(**self.__dict__)

    @property
    def inputs(self):
        return self.fields.get_inputs()

    @property
    def outputs(self):
        return self.fields.get_outputs()


class RegFields:
    def __init__(self, fields, regifdict):
        self.__dict__.update(regifdict)
        self._inputs = []
        self._outputs = []
        self.fn_len = self._get_fd_mxlen(fields)
        self.fd_align = self.fn_len + 10 if self.cdc else self.fn_len + 3
        self.fields = self._getfields(fields, regifdict)
        self.sync_list = self.sorts(("RC", "RWC")) if self.cdc else []
        self.rww_list = self.sorts(("RWW",))
        self.rwc_list = self.sorts(("RWC",))
        self.rwp_list = self.sorts(("RWP",))
        self.rwt_list = [fd for fd in self.sorts(("RWT",)) if not fd.wprotect]
        self.rc_list = self.sorts(("RC",))
        self.w_type = ("RW", "WO", "RWT", "RWC", "RWP", "RWW")
        self.cg_type = ("RW", "WO", "RWT")
        self.else_type = ("RWC", "RWP", "RWW")
        self.nmw_type = ("RWC", "RWP", "RWW") if (self.rw_cg) else ("RW", "WO", "RWC", "RWP", "RWW")
        self.rst_type = ("RWC", "RWP", "RWW") if (self.rw_cg) else ("RW", "WO", "RWT", "RWC", "RWP", "RWW")
        self.cg_rst_list = self.sorts(self.cg_type)
        self.cg_list = [fd for fd in self.sorts(("RW", "WO")) if not fd.wprotect]
        self.elselist = self.sorts(self.else_type)
        self.rst_list = self.sorts(self.rst_type)
        self.nocgw_list = [fd for fd in self.sorts(self.nmw_type) if not fd.wprotect]
        self.fds_with_reserved = self._getfields(fields, regifdict, 0);
        self.P, self.beP = self.get_protect_dicts(self.rst_type)
        self.cgP, self.cgbeP = self.get_protect_dicts(self.cg_type)

    def dynamic_check(self):
        msg = self._section_continue_check()
        msg += self._protect_check()
        for fd in self.fields:
            msg += fd.rst_val_check().format(**self.__dict__)
        return msg

    def _section_continue_check(self):
        bttps = []
        for fd in self.fds_with_reserved:
            bttps.insert(0, fd)
        msg = ""
        if bttps[0].bttp[0]:
            msg += "Error:The reg '{reg_name:{reg_align}}' first section {section} not begin with 0 \n".format(
                **bttps[0].__dict__)
        if bttps[-1].bttp[-1] != self.data_widthp1:
            msg += "Error:The reg '{reg_name:{reg_align}}' last section {section} not end with {dw}\n".format(
                dw=self.data_widthp1, **bttps[-1].__dict__)
        if len(bttps) > 1:
            for i in range(1, len(bttps)):
                if (bttps[i - 1].bttp[-1] + 1) != bttps[i].bttp[0]:
                    msg += "Error:The reg '{reg_name:{reg_align}}' section {s2} {section} not continuity \n".format(
                        s2=bttps[i].section, **bttps[i - 1].__dict__)
        return msg

    def _protect_check(self):
        msg = ""
        for section in self.beP:
            if section not in self.P:
                msg += "Error:The reg '{:{reg_align}}' field '{}' protect section {} not exists\n" \
                    .format(self.reg_name, self.beP[section][0].name, section, reg_align=self.reg_align)
            else:
                for bepfd in self.beP[section]:
                    if bepfd.bits_width != self.P[bepfd.wprotect].bits_width and self.P[bepfd.wprotect].type == "BP":
                        pfd = self.P[section]
                        msg += "Error:The reg '{:{reg_align}}' field '{}' width don't macth {} type '{}' width,can't protect bit by bit\n" \
                            .format(self.reg_name, bepfd.name, pfd.type, pfd.name, reg_align=self.reg_align)
        for section in self.P:
            pfd = self.P[section]
            if pfd.wprotect:
                msg += "Error:The reg '{:{reg_align}}' {} type field '{}' wprotect:'{}' not empty ,clear it\n" \
                    .format(self.reg_name, pfd.type, pfd.name, pfd.wprotect, reg_align=self.reg_align)
        return msg

    def _getfields(self, ofields, regifdict, eliminate=1):
        _fieldlist = []
        upperpara = regifdict.copy()
        upperpara["reg_name"] = self.reg_name
        upperpara["fd_align"] = self.fd_align
        for ofield in ofields:
            if not eliminate:
                _fieldlist.append(Field(ofield, upperpara))
            elif ofield[1] != "RESERVED":
                _fieldlist.append(Field(ofield, upperpara))
        return _fieldlist

    def _get_fd_mxlen(self, fields):
        _len = 0
        for field in fields:
            if len(field[1]) > _len:
                _len = len(field[1])
        return _len

    def get_protect_dicts(self, pt_type):
        P = {fd.section: fd for fd in self.fields if fd.type in ("CP", "BP")}
        beP = {}
        if not P:
            return P, beP
        for fd in self.fields:
            pt_section = fd.get_protect()
            if fd.type not in pt_type:
                continue
            if re.match(r"\[(\d+|\d+\:\d+)\]\Z", pt_section):
                if pt_section not in beP:
                    beP[pt_section] = [fd]
                else:
                    beP[pt_section].append(fd)
        return P, beP

    def sorts(self, types, exclude=False):
        lists = []
        for fd in self.fields:
            if exclude:
                if fd.type not in types:
                    lists.append(fd)
            else:
                if fd.type in types:
                    lists.append(fd)
        return lists

    def protect_w_L1(self, p_fd, bep_fds):
        p_section = p_fd.section
        w_valid_val = p_fd.get_protect()
        bprotects = ""
        cprotects = ""
        if p_fd.type == "BP":
            for fd in bep_fds:
                template = BP_RWT_L0 if fd.type == "RWT" else BP_L0
                for bit in fd.bits:
                    bitsel = "[{}]".format(bit) if len(fd.bits) > 1 else ""
                    bmibitsel = "[{}]".format(fd.bmi_bits[bit])
                    bmibitselp = "[{}]".format(p_fd.bmi_bits[bit])
                    if_name = fd.if_name
                    types = fd.type
                    bprotects = template.format(**locals()) + bprotects
            return "\n        //BP(bitwise protection) wirte" + bprotects
        elif p_fd.type == "CP":
            for fd in bep_fds:
                template = RWT_L1 if fd.type == "RWT" else W_L1
                bus_section = fd.section
                cprotects += fd.render(template, tab=4 * " ")  # tab
            return W_PROTECTE_L2.format(**locals())
        return ""

    def protect_w_L2(self, P, beP):
        _protect_part = ""
        for section in beP:
            _protect_part += self.protect_w_L1(P[section], beP[section])
        return _protect_part if beP else ""

    def merge(self, template, lists, **args):
        _merge = ""
        for fd in lists:
            _merge += fd.render(template, tab="", **args)
        return _merge if lists else ""

    def else_op(self):
        headrwc = "\n        //HW clear to reset value operation"
        headrww = "\n        //HW set operation"
        fieldrwp = self.merge(W_RST_L1, self.rwp_list)
        fieldrwc = self.merge(RWC_RST_L1, self.rwc_list, hw=self._hw_clc)
        fieldrww = self.merge(RWW_SET_L1, self.rww_list, seten=self._hw_set, setval=self._hw_setval)
        fieldrwc = headrwc + fieldrwc if fieldrwc else ""
        fieldrww = headrww + fieldrww if fieldrww else ""
        return ELSE_PART_L2.format(**locals()) if self.elselist else ""

    def w_always_L3(self):
        reg_name = self.reg_name
        field_rst = self.merge(W_RST_L1, self.rst_list)
        normal_wr = self.merge(W_L1, self.nocgw_list)
        rwt_wr = "" if (self.rw_cg) else (self.merge(RWT_L1, self.rwt_list) if self.rwt_list else "")
        protect_wr = self.protect_w_L2(self.P, self.beP)
        else_part = self.else_op()
        always_ff_begin = self.__dict__["always_ff_begin"]
        return W_ALWAYS_L3.format(**locals()) if self.rst_list else ""

    def w_cg_always_L3(self):
        reg_name = self.reg_name
        rw_rst = self.merge(W_RST_L1, self.cg_rst_list)
        rw_wr = self.merge(W_L1, self.cg_list)
        rwt_wr = self.merge(RWT_L1, self.rwt_list) if self.rwt_list else ""
        protect_wr = self.protect_w_L2(self.cgP, self.cgbeP)
        always_ff_begin = self.__dict__["always_ff_begin"]
        return W_CG_ALWAYS_L3.format(**locals()) if self.cg_rst_list else ""

    def fifo_filed_assign_L2(self):
        ret = ""
        for fd in self.fields:
            ret += fd.fifomap()
        return ret

    def fifo_push_declar(self):
        fifo_filed_assign = self.fifo_filed_assign_L2()
        render_dict = self.__dict__.copy()
        if_Comment = "FIFO interface"
        render_dict.update(locals())
        return FIFO_PUSH_L4.format(**render_dict)

    def cdc_sync_part_L2(self):
        cdc_sync_ffs = ""
        cdc_sync_rsts = ""
        cdc_sync_plss = ""
        for fd in self.sync_list:
            _hw = self._hw_clc if fd.type == "RWC" else self._hw_set
            _bswrap = fd.bwrap if fd.type == "RC" else ""
            cdc_sync_rsts += fd.render(CDC_RST_L1, hw=_hw)
            cdc_sync_ffs += fd.render(CDC_FF_L1, hw=_hw)
            cdc_sync_plss += fd.render(CDC_PLS_L1, hw=_hw, bswrap=_bswrap)
        return CDC_PART_L2.format(**locals()) if self.sync_list else ""

    def w_block(self):
        render_dict = self.__dict__.copy()
        w_always = self.w_always_L3()
        rw_cg_always = self.w_cg_always_L3() if (self.rw_cg) else ""
        rc_part = self.merge(RC_ALWYAS_L0, self.rc_list, hw=self._hw_set)
        cdc_part = self.cdc_sync_part_L2()
        if_Comment = "" if (rw_cg_always + w_always + rc_part + cdc_part) else "Read Only"
        render_dict.update(locals())
        return W_BLOCKS_L4.format(**render_dict)

    def read_express(self, bmi="bmi"):
        _read = ""
        for fd in self.fds_with_reserved:
            bitwrap = "[{}:0]".format(fd.bits_width - 1) if fd.bits_width > 1 else ""
            if (fd.type in ("WO", "BP", "CP")) or (fd.name == "RESERVED"):
                _read += "{}\'b0,".format(fd.bits_width)
            elif (fd.type == "RO"):
                _read += fd.name + bitwrap + ","
            else:
                _read += fd.if_name + bitwrap + ","
        _read = _read[0:-1]
        reg_name = self.reg_name
        reg_name_align = self.reg_align
        return R_L2.format(**locals())

    def declars(self):
        _io, _vdec, _wdec, _port = ("", "", "", "")
        # declars is an objects method ,so need initial class gobal vraible self._inputs,
        # self._outputs to prevent multiply add when the ojbects method call more than once
        self._inputs, self._outputs = [], []
        if (self.fifo_if):
            _io = "    output    {:11}{:{}} ,\n".format("", self.reg_name + "_push", self.gfd_align)
            self._outputs = ["{}".format(self.reg_name + "_push")]
            _port = "    .{name:{w}} ({name:{w}} ),//o\n".format(name=self.reg_name + "_push", w=self.gfd_align)
            return _io, _vdec, _wdec, _port
        for fd in self.fields:
            io, vdec, wdec, port = fd.declars()
            _io += io
            _vdec += vdec
            _wdec += wdec
            _port += port
            self._inputs += fd.get_inputs()
            self._outputs += fd.get_outputs()
        return _io, _vdec, _wdec, _port

    def get_inputs(self):
        return self._inputs

    def get_outputs(self):
        return self._outputs


class Field:
    def __init__(self, ofield, upperpara):
        self.__dict__.update(upperpara)
        self._inputs = []
        self._outputs = []
        self.section = ofield[0]
        self.name = ofield[1]
        self.type = ofield[2]
        self.wprotect = ofield[4]
        self.lock_sig = ofield[5] if self.type not in self.unlocktype else ""
        self.if_name = self.reg_pre + self.name
        self.oif_name = self.lock_pre + self.if_name
        self.bits_width = self._getbits()[0]
        self.bits = self._getbits()[2]
        self.bmi_bits = self._getbits()[1]
        self.bttp = self._getbits()[3]
        self.rstval = self._get_val(ofield[3])
        self.syncd = self._syncd if self.cdc else ""
        self.bwrap = "[{}:0]".format(self.bits_width - 1) if self.bits_width > 1 else ""

    def __str__(self):
        return "{}:{}{}".format(self.type, self.name, self.section)

    __repr__ = __str__

    def _get_val(self, value):
        strs = force2str(value, 1)
        if strs == "0":
            resetvalue = "{}'b{}".format(self.bits_width, strs)
        else:
            resetvalue = strs
        return resetvalue

    def rst_val_check(self):
        msg = ""
        if self.bits_width != int(self.rstval.split("'")[0]):
            msg += "Error:The reg '{{reg_name:{{reg_align}}}}' field {name:10} width '{bits_width}' don't macth it reset value '{rstval}' \n".format(
                **self.__dict__)
        decm = re.match("(\d+)'d(\d+)", self.rstval)
        binm = re.match("(\d+)'b([01]+)", self.rstval)
        hexm = re.match("(\d+)'h([\da-f]+)", self.rstval)
        max_val = 2 ** int(self.bits_width)
        if decm:
            rel_val = int(decm.group(2))
        elif binm:
            rel_val = int(binm.group(2), 2)
        elif hexm:
            rel_val = int(hexm.group(2), 16)
        else:
            raise "An except error, the reset value not match 'd 'b 'h 0 type"
        if rel_val > (max_val - 1):
            msg += "Error:The reg '{{reg_name:{{reg_align}}}}' field {name:10} reset value {rstval} exceed {section} supported max value' \n".format(
                **self.__dict__)
        return msg

    def _getbits(self):
        multi_bit = re.match(r"\[(?P<b>\d+):(?P<s>\d+)\]", self.section)
        single_bit = re.match(r"\[(?P<b>\d+)\]", self.section)
        if multi_bit:
            hsb = int(multi_bit.group(1))
            lsb = int(multi_bit.group(2))
            bits_len = hsb - lsb + 1
            bmi_bits = range(lsb, hsb + 1)
            bits = range(0, bits_len)
            bittuple = (lsb, hsb)
        elif single_bit:
            bits_len = 1
            bmi_bits = (int(single_bit.group(1)),)
            bits = (0,)
            bittuple = bmi_bits
        else:
            raise "bit section error"
        return (bits_len, bmi_bits, bits, bittuple)

    def get_protect(self):
        return self.rstval if self.type in ("BP", "CP") else force2str(self.wprotect)

    inputs = lambda self: self._render2(INPUT, bw=self.bwrap, pn=self.name, **self.__dict__)
    inputs_set = lambda self: self._render2(INPUT, bw="", pn=self.name + self._hw_set, **self.__dict__)
    inputs_setval = lambda self: self._render2(INPUT, bw=self.bwrap, pn=self.name + self._hw_setval, **self.__dict__)
    inputs_sets = lambda self: self._render2(INPUT, bw=self.bwrap, pn=self.name + self._hw_set, **self.__dict__)
    inputs_clc = lambda self: self._render2(INPUT, bw="", pn=self.name + self._hw_clc, **self.__dict__)
    outputs = lambda self: self._render2(OUTPUT, bw=self.bwrap, pn=self.if_name, **self.__dict__)
    inport = lambda self: self._render2(IPORTS, pn=self.name, pn2=self.name + self.bwrap, **self.__dict__)
    inport_set = lambda self: self._render2(IPORTS, pn=self.name + self._hw_set, pn2=self.name + self._hw_set,
                                            **self.__dict__)
    inport_setval = lambda self: self._render2(IPORTS, pn=self.name + self._hw_setval,
                                               pn2=self.name + self._hw_setval + self.bwrap, **self.__dict__)
    inport_sets = lambda self: self._render2(IPORTS, pn=self.name + self._hw_set,
                                             pn2=self.name + self._hw_set + self.bwrap, **self.__dict__)
    inport_clc = lambda self: self._render2(IPORTS, pn=self.name + self._hw_clc, pn2=self.name + self._hw_clc,
                                            **self.__dict__)
    outport = lambda self: self._render2(OPORTS, pn2=self.if_name + self.bwrap, **self.__dict__)
    outport_lock = lambda self: self._render2(OPORTS, pn2=self.oif_name + self.bwrap, **self.__dict__)
    reg_if = lambda self: self._render2(REG, bw=self.bwrap, pn=self.if_name, **self.__dict__)
    wire_if = lambda self: self._render2(WIRE, bw=self.bwrap, pn=self.if_name, **self.__dict__)
    wire_oif = lambda self: self._render2(WIRE, bw=self.bwrap, pn=self.oif_name, **self.__dict__)
    reg_clc_q1 = lambda self: self._render2(REG, bw="", pn=self.name + self._hw_clc + "_q1", **self.__dict__)
    reg_clc_q2 = lambda self: self._render2(REG, bw="", pn=self.name + self._hw_clc + "_q2", **self.__dict__)
    wire_clc_syncd = lambda self: self._render2(WIRE, bw="", pn=self.name + self._hw_clc + self.syncd, **self.__dict__)
    reg_set_mq1 = lambda self: self._render2(REG, bw=self.bwrap, pn=self.name + self._hw_set + "_q1", **self.__dict__)
    reg_set_mq2 = lambda self: self._render2(REG, bw=self.bwrap, pn=self.name + self._hw_set + "_q2", **self.__dict__)
    wire_set_msyncd = lambda self: self._render2(WIRE, bw=self.bwrap, pn=self.name + self._hw_set + self.syncd,
                                                 **self.__dict__)
    fifo_map = lambda self: self._render2(FIFOMAP, pn=self.if_name, **self.__dict__)

    def fifomap(self):
        if self.type == "WO":
            return self.fifo_map()

    def declars(self):
        "inital to prevent mutiply add when method decalrs be called more than once "
        self._inputs, self._outputs = [], []
        _io = ""
        _vdec = ""
        _wdec = ""
        _port = ""
        '''
               RW RC BP CP RO WO RWW RWC RWT RWP  |remove BP,CP
               RO,RC,RWC,RWW,WO,RW,RWT,RWP        |dived to 4 group
        case0 (RO)               |
        case1 (RWC)              |CDC,LOCK
        case2 (RC)               |CDC
        case3 (RWW),             |LOCK
        case4 (WO, RW, RWT, RWP) |LOCK
        '''
        if self.type not in ("BP", "CP"):
            "case0 (RO)"
            if self.type == "RO":
                _io += self.inputs()
                _port += self.inport()
                return _io, _vdec, _wdec, _port
            "case1 (RWC),             |CDC,LOCK"
            if self.type == "RWC":
                if self.cdc:
                    _vdec += self.reg_clc_q1()
                    _vdec += self.reg_clc_q2()
                    _vdec += self.wire_clc_syncd()
                _io += self.inputs_clc()
                _port += self.inport_clc()
                _port += self.outport()
                _io += self.outputs()
                _vdec += self.reg_if()
                if self.lock_sig:
                    _wdec += self.reg_if()
                    _wdec += self.wire_oif()
                else:
                    _wdec += self.wire_if()
                return _io, _vdec, _wdec, _port
            "case3 (RC),             |CDC"
            if self.type == "RC":
                if self.cdc:
                    _vdec += self.reg_set_mq1()
                    _vdec += self.reg_set_mq2()
                    _vdec += self.wire_set_msyncd()
                _io += self.inputs_sets()
                _port += self.inport_sets()
                _port += self.outport()
                _io += self.outputs()
                _vdec += self.reg_if()
                return _io, _vdec, _wdec, _port
            "case4 (RWW),             |LOCK"
            if self.type == "RWW":
                _io += self.inputs_set()
                _io += self.inputs_setval()
                _port += self.inport_set()
                _port += self.inport_setval()
                _io += self.outputs()
                _vdec += self.reg_if()
                if self.lock_sig:
                    _wdec += self.reg_if()
                    _wdec += self.wire_oif()
                    _port += self.outport_lock()
                else:
                    _wdec += self.wire_if()
                    _port += self.outport()
                return _io, _vdec, _wdec, _port
            "case5 (WO, RW, RWT, RWP) |LOCK"
            _io += self.outputs()
            _vdec += self.reg_if()
            if self.lock_sig:
                _wdec += self.reg_if()
                _wdec += self.wire_oif()
                _port += self.outport_lock()
            else:
                _wdec += self.wire_if()
                _port += self.outport()
        return _io, _vdec, _wdec, _port

    def _render2(self, template, **dargs):
        if template == INPUT:
            self._inputs.append(dargs["pn"])  # Be careful! ,class varible be changed in function
        elif template == OUTPUT:
            self._outputs.append(dargs["pn"])  # the function need be class private method start with "_"
        return template.format(**dargs)

    def get_inputs(self):
        return self._inputs

    def get_outputs(self):
        return self._outputs

    def render(self, template, **dargs):
        render_dict = self.__dict__.copy()
        render_dict.update(dargs)
        bitsel = ""
        render_dict.update(locals())
        if self.type == "RC":
            _rc = ""
            for bit in self.bits:
                bitsel = "[{}]".format(bit) if len(self.bits) > 1 else ""
                bmibitsel = "[{}]".format(self.bmi_bits[bit])
                render_dict.update(locals())
                _rc = template.format(**render_dict) + _rc
            return _rc
        return template.format(**render_dict)


def static_check(regs):
    msg = ""
    line = 0
    regs_unq_msg, addrs_unq_msg, fields_unq_msg = ("",) * 3
    regs_unq, addrs_unq, fields_unq = ([],) * 3
    for reg in regs:
        if reg[1] in addrs_unq:
            addrs_unq_msg += "Error: position {:3}{}: Regname {:20} already exist\n".format(line + 10, chr(65 + 1),
                                                                                            reg[1])
        if reg[0] in regs_unq:
            regs_unq_msg += "Error: position {:3}{}: Addrs   {:20} already exist\n".format(line + 10, chr(65 + 0),
                                                                                           reg[0])
        addrs_unq.append(reg[1])
        regs_unq.append(reg[0])
        for i in range(len(reg_format)):
            strs = force2str(reg[i])
            msg += format_check(strs, line, i, reg_format[i])
        for field in reg[4]:
            if field[1] in fields_unq and field[1] != "RESERVED":
                fields_unq_msg += "Error: position {:3}{}: Field   {:20} already exist\n".format(line + 10, chr(65 + 4),
                                                                                                 field[1])
            fields_unq.append(field[1])
            for i in range(len(field_format)):
                strs = force2str(field[i])
                msg += format_check(strs, line, i + 4, field_format[i])
            line += 1
    return msg + regs_unq_msg + addrs_unq_msg + fields_unq_msg


def force2str(value, ignorecase=0):
    _tmp = str(int(value)) if type(value) == float else \
        str(value) if type(value) == int else value
    return _tmp.lower() if ignorecase else _tmp


def format_check(strs, line, row, info):
    if not re.match(info[1], strs):
        wstrs = '"' + strs + '"'
        return "Error: position {:3}{}: {:28} format error! Suggest like : {}\n" \
            .format(line + 10, chr(65 + row), wstrs, info[2])
    else:
        return ""


def title_check(bs):
    msg = ""
    for formats in other_format:
        line = formats[0]
        row = formats[1]
        info = formats[2:]
        msg += format_check(force2str(bs.cell_value(line, row)), line, row, info)
    return msg


def split2fixwidth(strlists, width):
    _out = "\n    "
    _len = 0
    for _str in strlists:
        _step = (len(_str) + 2)
        _len += _step
        if _len > width and (width - (_len - _step)) < 8:
            _len = 0
            _out += ("\n    " + _str + ", ")
        else:
            _out += (_str + ", ")
    return _out

def ispow2(x):
    return False if x & (x-1) else True

def groupSize(info):
    if info:
        gs = int(info)
        if(ispow2(gs)):
            return gs
        else:
            print("Error! %d is not pow2, 32, 64, 128 is recommended, Fix at '5K' postion, then re-run again\n"%gs)
            exit()
    else:
        return 0

class Regif:
    def __init__(self, bs):
        self.regs = self.extract_regs(bs)
        self._static_check(bs)
        self.modulename = re.match(other_format[0][3], bs.cell(2, 1).value).group()
        self.addr_width = force2str(bs.cell(3, 5).value)
        self.rw_cg = True if (force2str(bs.cell(3, 10).value).upper() == "YES") else False
        self.data_width = force2str(bs.cell(4, 5).value)
        self.reg_pre = force2str(bs.cell(6, 5).value).strip('"')
        self.group_size = groupSize(bs.cell(4, 10).value)
        self.unlocktype = ("RO", "RC", "RWT", "RWP", "CP", "BP")
        self.lock_pre = "o"
        self.page_width = 50
        self._syncd = "_syncd"
        self._hw_clc = "_hw_clc"
        self._hw_set = "_hw_set"
        self._hw_setval = "_hw_setval"
        self.cdc = 0 if bs.cell(5, 1).value == "SYNC" else 1
        self.reg_align = self.get_reg_name_mxlen()
        self.gfd_align = self.get_fd_name_mxlen() + 11 if self.cdc else self.get_fd_name_mxlen() + 9
        self.addr_widthp1 = int(self.addr_width) - 1
        self.data_widthp1 = int(self.data_width) - 1
        self.always_ff_begin = ALWAYS_BEGIN_MACRO if(bs.cell(5,10).value.upper() == "YES") else ALWAYS_BEGIN
        self.macros = RESET_MACRO if(bs.cell(5,10).value.upper()=="YES") else ""
        self.reg_objs = self.gen_reg_objects()
        self._dynamic_check()

    def extract_regs(self, bs):
        regs = []
        nrows = bs.nrows
        ncols = bs.ncols
        address_idx = 0
        for i in range(9, nrows):
            row_data = bs.row_values(i)
            if row_data[0] != u'':
                regs.append(row_data[0:4])
                regs[address_idx].append([])
                regs[address_idx][4].append(row_data[4:ncols])
                address_idx += 1
            else:
                regs[address_idx - 1][4].append(row_data[4:ncols])
        return regs

    def _static_check(self, bs):
        print("Static check ..... ")
        msg = title_check(bs)
        msg += static_check(self.regs)
        if msg:
            msg += "\nFixed all those format Error first, then regenerate again !\n"
            print(msg)
            exit(0)
        else:
            print("Static check Pass! \n")
        return

    def _dynamic_check(self):
        print("Dynamic check ..... ")
        msg = ""
        for reg_obj in self.reg_objs:
            msg += reg_obj.dynamic_check()
        if msg:
            msg += "\nFixed all those format Error first, then regenerate again !\n"
            print(msg)
            exit(0)
        else:
            print("Dynamic check Pass! \n")
        return

    def _fileds_correlation_check(self):
        msg = ""
        for reg in self.reg_objs:
            msg += reg.dynamic_check()
        return msg

    def gen_reg_objects(self):
        dictcopy = self.__dict__.copy()
        dictcopy.pop("regs")
        reg_objects = [Reg(reg, dictcopy) for reg in self.regs if not Reg(reg, dictcopy).rtl_free]
        return reg_objects

    def get_fd_name_mxlen(self):
        _len = 0
        for reg in self.regs:
            for fd in reg[4]:
                if len(fd[1]) > _len:
                    _len = len(fd[1])
        return _len

    def get_reg_name_mxlen(self):
        _len = 0
        for reg in self.regs:
            if len(reg[1]) > _len:
                _len = len(reg[1])
        return _len

    def getlockdict(self):
        _lockdict = {}
        for reg in self.reg_objs:
            for fd in reg.fields.fields:
                if (fd.lock_sig and fd.type not in self.unlocktype):
                    if fd.lock_sig not in _lockdict:
                        _lockdict[fd.lock_sig] = [fd]
                    elif not reg.fifo_if:
                        _lockdict[fd.lock_sig].append(fd)
        return _lockdict

    def gen_lock_segment(self):
        lockdict = self.getlockdict()
        _lock_always = ""
        _lock_sig = ""
        for locksig in lockdict:
            lock_rsts = ""
            lock_express = ""
            lock_en = locksig
            for fd in lockdict[locksig]:
                lock_rsts += fd.render(W_RST_L1, fd_align=self.gfd_align - 3)
                lock_express += fd.render(LOCK_EXPRESS_L1, fd_align=self.gfd_align - 3)
            _lock_always += LOCK_ALWAYS_L2.format(**locals())
            _lock_sig += INPUT.format(bw="", pn=locksig, **fd.__dict__)
        return _lock_sig, _lock_always

    def gen_segment(self):
        write_part, iodeclars, vdeclars, wdeclars, instports, defines, undefines, reads = ("",) * 8
        gs = self.group_size
        read_part = ""
        if(gs != 0):
            grpreads = ""
            for i in range(0, len(self.reg_objs), gs):
                reads  = ""
                grpid  = int(i/gs)
                bmi = "grp%d" % grpid
                for reg in self.reg_objs[i : i+gs]:
                    reads += reg.read(bmi)
                read_part += R_ALWAYS.format(reads=reads, bmi=bmi, **self.__dict__)
                if(i == 0):
                    grpreads  += RG_RD.format(i = grpid)
                else:
                    grpreads  += RG_RDE.format(i = grpid)
            read_part += RG_ALWAYS.format(grpreads=grpreads, bmi=bmi, **self.__dict__ )
            print("Info! bmi_rdvld have 2 cycle delay when bmi_rd \n")
        else:
            reads = ""
            for reg in self.reg_objs:
                reads += reg.read()
            read_part = R_ALWAYS.format(reads=reads, bmi="bmi", **self.__dict__)

        for reg in self.reg_objs:
            write_part += reg.write
            iodeclars += reg.iodeclars
            vdeclars += reg.vdeclars
            wdeclars += reg.wdeclars
            instports += reg.instports
            defines += reg.defines
            undefines += reg.undefines
        cg_io = CG_IO if (self.rw_cg) else ""
        bmi_io = BMI_IO.format(cg_io, addr_widthp1=int(self.addr_width) - 1, data_widthp1=int(self.data_width) - 1)
        instports = instports[:-5] + " " + instports[-4:-1]
        return bmi_io, write_part, iodeclars, vdeclars, wdeclars, instports, defines, undefines, read_part

    def get_ports(self):
        inputs, outputs = [], []
        for reg in self.reg_objs:
            inputs += reg.inputs
            outputs += reg.outputs
        return inputs, outputs

    def genverilog(self):
        bmi_io, write_part, iodeclars, vdeclars, wdeclars, instports, defines, undefines, read_part = self.gen_segment()
        macros = self.macros
        lock_io, lock_part = self.gen_lock_segment()
        inputs, outputs = self.get_ports()
        inputs_add_lock = inputs + ["hw_clk", "hw_rstn"] + list(self.getlockdict().keys())
        cgclk = "bmi_clk_wr_cg, " if self.rw_cg else ""
        cg_port = CG_PORT if self.rw_cg else ""
        vinparts = split2fixwidth(inputs, self.page_width)
        winparts = split2fixwidth(inputs_add_lock, self.page_width)
        # outparts = split2fixwidth(outputs, self.page_width)[:-2]  # remove last ",\n"
        def outparts():
            if(len(outputs) == 0):
                return ""
            else:
                return "," + split2fixwidth(outputs, self.page_width)[:-2]  # remove last ",\n"
        vioports = DEC_PORTS.format(cgclk, vinparts, outparts())
        wioports = DEC_PORTS.format(cgclk, winparts, outparts())
        addr_widthp1 = int(self.addr_width) - 1
        data_widthp1 = int(self.data_width) - 1
        modulename = self.modulename
        v_file = VFILE.format(**locals())
        w_file = WFILE.format(**locals())
        fp1 = open(modulename + "_regif.v", "w")
        fp1.write(v_file)
        fp1.close()
        print(modulename + "_regif.v      generate done!")
        if lock_io:
            fp2 = open(modulename + "_regif_wrap.v", "w")
            fp2.write(w_file)
            fp2.close()
            print(modulename + "_regif_wrap.v generate done!\n")
        # ----------------------------------------------------------------


# xxx_Register_Manual.docx  generate
# ----------------------------------------------------------------

from docx.shared import Pt
from docx.shared import Cm
from docx.shared import RGBColor
from docx.shared import Length
from docx.oxml.ns import qn


def creat_reg_table(document):
    table_style = document.styles["Table Grid"]
    table_style.font.size = Pt(10.5)
    table = document.add_table(rows=1, cols=0, style=table_style)
    aval_width = Length(document._block_width).cm
    table.add_column(Cm(aval_width * 0.1))
    table.add_column(Cm(aval_width * 0.15))
    table.add_column(Cm(aval_width * 0.1))
    table.add_column(Cm(aval_width * 0.15))
    table.add_column(Cm(aval_width * 0.45))
    table.cell(0, 0).text = u"位数"
    table.cell(0, 1).text = u"名称"
    table.cell(0, 2).text = u"方向"
    table.cell(0, 3).text = u"复位值"
    table.cell(0, 4).text = u"描述"
    return table


def fill_reg_fields(table, fields):
    for field in fields:
        row_cells = table.add_row().cells
        reserved = True if field[1] == "RESERVED" else False
        row_cells[0].text = force2str(field[0])[1:-1]
        row_cells[1].text = force2str(field[1])
        row_cells[2].text = "N/A" if reserved else force2str(field[2])
        row_cells[3].text = "N/A" if reserved else force2str(field[3])
        row_cells[4].text = force2str(field[6])


def creat_regs_detail(document, regs):
    for reg in regs:
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.add_heading(u'{} ({})'.format(reg[1], reg[0]), 2)
        document.add_paragraph(u"偏移地址: {}".format(reg[0]))
        document.add_paragraph(u"说    明: {}".format(reg[2]))
        new_reg_table = creat_reg_table(document)
        fill_reg_fields(new_reg_table, reg[4])
    return


def gen_regif_doc(name, regs):
    document = Document()
    document.add_heading(u'Register Interface', 1)
    # craet_addrs_map(regs)
    creat_regs_detail(document, regs)
    document.save(u"{}_Register_Manual.docx".format(name))


def regif_args():
    parser = argparse.ArgumentParser(prog="regif",
                                     description='\033[1;34mRegister Interface tool(auto generator .v/.docx)\033[0m')

    parser.add_argument('excel', metavar='xxx_regif.xls', nargs='?', help='parser excel then generator .v/.docx')
    parser.add_argument("--init", metavar="xxx", help="creat xxx_regif.xls template at ./")
    args = parser.parse_args()
    return parser


def copyTemplate(name):
    regif = os.path.abspath(__file__)
    binpath = os.path.dirname(regif)
    template = os.path.join(binpath, "../regif/regif_template.xls")
    dest = "%s_regif.xls" % name
    if os.path.exists(dest):
        print("\n%s already exists !!!\n" % dest)
        exit(0)
    os.system("cp %s %s" % (template, dest))
    print("\nRegIf excel template %s created\n" % dest)


if __name__ == "__main__":
    parser = regif_args()

    args = parser.parse_args()

    if (args.init is not None):
        copyTemplate(args.init)
    elif (args.excel is not None):
        try:
            if (not os.path.exists(args.excel)):
                print("%s not found, check please, 'regif -h' for help" % args.excel)
                exit(0)
            wb = xlrd.open_workbook(args.excel)
            bs = wb.sheet_by_index(0)
        except:
            raise
        regif = Regif(bs)
        regif.genverilog()
        modulename = regif.modulename
        try:
            gen_regif_doc(modulename, regif.regs)
        except IOError as e:
            print(e)
            print(u"Make sure the {}_Register_Manual.docx is not opened first".format(modulename))
            exit(0)
        print("{}_Register_Manual.docx generate done!\n".format(modulename))
    else:
        parser.print_help()
